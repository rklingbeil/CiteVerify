"""Create test briefs for CiteVerify accuracy testing.

Brief A: All real cases with accurate quotes and characterizations
Brief B: Real cases with fabricated/altered quotes and characterizations
Brief C: Mix of real and completely hallucinated cases

These briefs test the full CiteVerify pipeline end-to-end.
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(12)
    return p


def _add_block_quote(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.right_indent = Inches(0.5)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    return p


# ─────────────────────────────────────────────────────────────────────────
# BRIEF A: All Real Cases — Accurate Quotes & Characterizations
# Expected: All citations verified, all quotes accurate
# ─────────────────────────────────────────────────────────────────────────

def create_brief_a():
    doc = Document()

    _add_heading(doc, "MEMORANDUM IN SUPPORT OF MOTION FOR SUMMARY JUDGMENT")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Johnson v. Pacific Northwest Industries, Inc.")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # spacer

    _add_heading(doc, "I. STANDARD OF REVIEW", level=2)

    _add_para(doc,
        'Summary judgment is appropriate when "there is no genuine dispute as to any '
        'material fact and the movant is entitled to judgment as a matter of law." '
        'Fed. R. Civ. P. 56(a). The Supreme Court has held that the moving party bears '
        'the initial burden of demonstrating the absence of a genuine issue of material '
        'fact. Celotex Corp. v. Catrett, 477 U.S. 317, 323 (1986).'
    )

    _add_para(doc,
        'In evaluating a motion for summary judgment, courts must view the evidence '
        'in the light most favorable to the nonmoving party. Anderson v. Liberty Lobby, '
        'Inc., 477 U.S. 242, 255 (1986). The Court explained:'
    )

    _add_block_quote(doc,
        '"The evidence of the non-movant is to be believed, and all justifiable '
        'inferences are to be drawn in his favor."'
    )

    _add_para(doc,
        'Anderson v. Liberty Lobby, Inc., 477 U.S. at 255. However, the nonmoving '
        'party must present more than a "mere scintilla" of evidence to survive summary '
        'judgment. Id. at 252.'
    )

    _add_heading(doc, "II. THE FIRST AMENDMENT PROTECTS COMMERCIAL SPEECH", level=2)

    _add_para(doc,
        'The First Amendment, as applied through the Fourteenth Amendment, protects '
        'commercial speech. Virginia State Board of Pharmacy v. Virginia Citizens '
        'Consumer Council, Inc., 425 U.S. 748 (1976). The Supreme Court established '
        'a four-part test for determining when commercial speech regulations are '
        'permissible in Central Hudson Gas & Electric Corp. v. Public Service Commission '
        'of New York, 447 U.S. 557 (1980). Under Central Hudson, the government may '
        'restrict commercial speech only if: (1) the speech concerns lawful activity and '
        'is not misleading; (2) the government interest is substantial; (3) the regulation '
        'directly advances the government interest; and (4) the regulation is no more '
        'extensive than necessary. Id. at 566.'
    )

    _add_heading(doc, "III. DUE PROCESS REQUIRES NOTICE AND OPPORTUNITY TO BE HEARD", level=2)

    _add_para(doc,
        'The Fourteenth Amendment provides that no state shall "deprive any person of '
        'life, liberty, or property, without due process of law." The Supreme Court in '
        'Mathews v. Eldridge, 424 U.S. 319 (1976), established a three-factor balancing '
        'test for determining what process is due:'
    )

    _add_block_quote(doc,
        '"First, the private interest that will be affected by the official action; '
        'second, the risk of an erroneous deprivation of such interest through the '
        'procedures used, and the probable value, if any, of additional or substitute '
        'procedural safeguards; and finally, the Government\'s interest, including the '
        'function involved and the fiscal and administrative burdens that the additional '
        'or substitute procedural requirement would entail."'
    )

    _add_para(doc,
        'Mathews v. Eldridge, 424 U.S. at 335. The Ninth Circuit has applied the '
        'Mathews balancing test in numerous contexts. See Goldberg v. Kelly, 397 U.S. '
        '254 (1970) (holding that welfare recipients must receive an evidentiary hearing '
        'before termination of benefits).'
    )

    _add_heading(doc, "IV. QUALIFIED IMMUNITY ANALYSIS", level=2)

    _add_para(doc,
        'Government officials performing discretionary functions are generally shielded '
        'from liability for civil damages insofar as their conduct does not violate '
        '"clearly established" statutory or constitutional rights. Harlow v. Fitzgerald, '
        '457 U.S. 800, 818 (1982). The Supreme Court refined this analysis in Ashcroft '
        'v. al-Kidd, 563 U.S. 731 (2011), holding that qualified immunity protects '
        '"all but the plainly incompetent or those who knowingly violate the law." '
        'Id. at 743.'
    )

    _add_para(doc,
        'The qualified immunity inquiry involves two prongs: (1) whether the facts '
        'alleged make out a violation of a constitutional right, and (2) whether the '
        'right at issue was "clearly established" at the time of defendant\'s alleged '
        'misconduct. Pearson v. Callahan, 555 U.S. 223, 232 (2009). Courts may address '
        'either prong first. Id. at 236.'
    )

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'For the foregoing reasons, Plaintiff respectfully requests that this Court '
        'grant summary judgment in her favor.'
    )

    path = os.path.join(OUTPUT_DIR, "brief_a_all_real.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────
# BRIEF B: Real Cases — Fabricated Quotes & Altered Characterizations
# Expected: Citations verified, but quotes/characterizations flagged
# ─────────────────────────────────────────────────────────────────────────

def create_brief_b():
    doc = Document()

    _add_heading(doc, "MEMORANDUM IN OPPOSITION TO MOTION TO DISMISS")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Reynolds v. National Healthcare Corp.")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    _add_heading(doc, "I. STANDARD FOR MOTION TO DISMISS", level=2)

    # Real case, FABRICATED quote (this is NOT what Twombly actually says)
    _add_para(doc,
        'To survive a motion to dismiss under Rule 12(b)(6), a complaint must contain '
        '"sufficient factual matter, accepted as true, to state a claim that is '
        'plausible beyond a reasonable doubt." Bell Atlantic Corp. v. Twombly, '
        '550 U.S. 544, 570 (2007). The Court emphasized that this standard requires '
        'more than mere labels and conclusions.'
    )

    # Real case, FABRICATED quote (Iqbal doesn't say this)
    _add_para(doc,
        'The Supreme Court clarified this standard in Ashcroft v. Iqbal, 556 U.S. 662 '
        '(2009), holding that:'
    )

    _add_block_quote(doc,
        '"A claim has facial plausibility when the plaintiff pleads factual content '
        'that allows the court to draw the automatic inference that the defendant '
        'is liable for the misconduct alleged. The plausibility standard requires '
        'substantially more than a sheer possibility."'
    )

    _add_para(doc,
        'Ashcroft v. Iqbal, 556 U.S. at 678.'
    )

    _add_heading(doc, "II. DEFENDANT OWED A DUTY OF CARE", level=2)

    # Real case, MISLEADING characterization (Palsgraf is about proximate cause/foreseeability,
    # not about "absolute duty to prevent all foreseeable harm")
    _add_para(doc,
        'It is well established that a defendant owes an absolute duty to prevent all '
        'foreseeable harm to any person within the zone of danger. Palsgraf v. Long '
        'Island Railroad Co., 248 N.Y. 339 (1928). Under Palsgraf, the question of '
        'negligence is always resolved in favor of the plaintiff when any risk of harm '
        'was foreseeable.'
    )

    # Real case, FABRICATED holding (Miranda is about custodial interrogation warnings,
    # not about "all government interactions")
    _add_para(doc,
        'Similarly, the right to counsel attaches at all government interactions where '
        'a person\'s liberty is at stake. Miranda v. Arizona, 384 U.S. 436 (1966). The '
        'Court held that "the prosecution may not use statements stemming from any '
        'governmental interaction unless it demonstrates the use of procedural '
        'safeguards effective to secure the privilege against self-incrimination." '
        'Id. at 444.'
    )

    _add_heading(doc, "III. THE ECONOMIC LOSS DOCTRINE DOES NOT APPLY", level=2)

    # Real case, ALTERED quote (the actual quote from East River uses different language)
    _add_para(doc,
        'The economic loss doctrine bars recovery in tort for purely economic losses '
        'arising from breach of contract. East River Steamship Corp. v. Transamerica '
        'Delaval, Inc., 476 U.S. 858 (1986). The Court stated that "a manufacturer '
        'in a commercial relationship has no duty in tort to prevent purely financial '
        'disappointment of a purchaser\'s economic expectations." Id. at 871.'
    )

    # Real case with CORRECT characterization (for comparison)
    _add_para(doc,
        'However, when physical injury or property damage accompanies the economic '
        'loss, tort recovery is available. See Saratoga Fishing Co. v. J.M. Martinac '
        '& Co., 520 U.S. 875 (1997) (holding that commercial fishermen could recover '
        'in tort for damage to additional equipment installed after purchase, as these '
        'items were "other property" not subject to the economic loss doctrine).'
    )

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'For the foregoing reasons, Defendant\'s Motion to Dismiss should be denied.'
    )

    path = os.path.join(OUTPUT_DIR, "brief_b_fabricated_quotes.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────
# BRIEF C: Mix of Real and Completely Hallucinated Cases
# Expected: Real cases verified; fake cases flagged as errors/unverifiable
# ─────────────────────────────────────────────────────────────────────────

def create_brief_c():
    doc = Document()

    _add_heading(doc, "PLAINTIFF'S TRIAL BRIEF")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Thompson v. Western Regional Medical Center")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    _add_heading(doc, "I. STATEMENT OF THE CASE", level=2)

    _add_para(doc,
        'Plaintiff brings this action for medical malpractice and negligent '
        'credentialing. The applicable standard of care is well established.'
    )

    _add_heading(doc, "II. APPLICABLE LEGAL STANDARDS", level=2)

    # REAL case — verifiable
    _add_para(doc,
        'Medical malpractice claims require proof of duty, breach, causation, and '
        'damages. The Supreme Court addressed the standard for expert testimony in '
        'Daubert v. Merrell Dow Pharmaceuticals, Inc., 509 U.S. 579 (1993), '
        'establishing that trial courts serve as gatekeepers for the admissibility '
        'of expert scientific testimony under Federal Rule of Evidence 702.'
    )

    # HALLUCINATED case — does not exist
    _add_para(doc,
        'The Ninth Circuit extended Daubert to medical malpractice cases in '
        'Whitfield v. Pacific Medical Associates, 387 F.3d 1042 (9th Cir. 2004), '
        'holding that "in medical malpractice actions, the trial court must ensure '
        'that expert testimony regarding standard of care is grounded in accepted '
        'medical methodology." Id. at 1048.'
    )

    # REAL case — verifiable
    _add_para(doc,
        'Expert testimony must be both relevant and reliable. Kumho Tire Co. v. '
        'Carmichael, 526 U.S. 137 (1999). The Court in Kumho extended Daubert\'s '
        'gatekeeping function to all expert testimony, not just scientific testimony.'
    )

    _add_heading(doc, "III. NEGLIGENT CREDENTIALING", level=2)

    # HALLUCINATED case — does not exist
    _add_para(doc,
        'Hospitals have an independent duty to verify the credentials of physicians '
        'granted privileges. In Morrison v. St. Luke\'s Regional Medical Center, '
        '498 F.3d 835 (8th Cir. 2007), the Eighth Circuit held that:'
    )

    _add_block_quote(doc,
        '"A hospital that grants privileges to a physician without conducting a '
        'reasonable investigation into that physician\'s competence and qualifications '
        'may be held directly liable for injuries resulting from that physician\'s '
        'negligent treatment of patients."'
    )

    _add_para(doc,
        'Morrison v. St. Luke\'s Regional Medical Center, 498 F.3d at 842.'
    )

    # REAL case — verifiable
    _add_para(doc,
        'The doctrine of corporate negligence holds hospitals to a duty of care in '
        'the selection and retention of medical staff. This principle was established '
        'in Darling v. Charleston Community Memorial Hospital, 33 Ill. 2d 326 (1965), '
        'which recognized that a hospital owes a direct duty to its patients to ensure '
        'the competency of its medical staff.'
    )

    _add_heading(doc, "IV. CAUSATION", level=2)

    # HALLUCINATED case with plausible-sounding citation
    _add_para(doc,
        'Under the "substantial factor" test for causation, the plaintiff need not '
        'prove that the defendant\'s conduct was the sole cause of injury. Rodriguez '
        'v. Northwestern Memorial Healthcare, 612 F.3d 544 (7th Cir. 2010). The '
        'Seventh Circuit explained that "where multiple acts of negligence combine to '
        'produce a single indivisible injury, each negligent actor is liable for the '
        'entire harm." Id. at 551.'
    )

    # REAL case — verifiable
    _add_para(doc,
        'The "but for" test remains the default standard for causation in most '
        'jurisdictions. See Burrage v. United States, 571 U.S. 204 (2014) (holding '
        'that "but for" causation is the default rule unless the statute provides '
        'otherwise).'
    )

    # HALLUCINATED case — plausible state court citation
    _add_para(doc,
        'Oregon courts have adopted a modified substantial factor test. See Kellerman '
        'v. Providence Health Systems, 348 Or 456 (2010) (holding that "a plaintiff '
        'in a medical malpractice action need only demonstrate that the defendant\'s '
        'negligence was a substantial factor in bringing about the harm, not that it '
        'was the predominant cause"). See also Chen v. Legacy Health Corp., 285 Or '
        'App 312 (2017).'
    )

    _add_heading(doc, "V. DAMAGES", level=2)

    # REAL case — verifiable
    _add_para(doc,
        'Compensatory damages in medical malpractice cases include both economic and '
        'noneconomic losses. The Supreme Court has recognized that due process limits '
        'punitive damage awards. BMW of North America, Inc. v. Gore, 517 U.S. 559 '
        '(1996). The Court established three guideposts for evaluating the '
        'constitutionality of punitive damages: (1) the degree of reprehensibility, '
        '(2) the ratio between punitive and compensatory damages, and (3) the '
        'difference between the punitive award and comparable civil penalties.'
    )

    # REAL case with FABRICATED specific ratio (State Farm says "few awards exceeding
    # single-digit ratio" will satisfy due process, not "3:1")
    _add_para(doc,
        'The Court later clarified that punitive damages generally cannot exceed a '
        '3:1 ratio to compensatory damages. State Farm Mutual Automobile Insurance '
        'Co. v. Campbell, 538 U.S. 408 (2003). The Court stated that "in practice, '
        'few awards exceeding a 3-to-1 ratio between punitive and compensatory '
        'damages will satisfy due process." Id. at 425.'
    )

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'Plaintiff respectfully requests that this Court enter judgment in her favor '
        'and award compensatory and punitive damages as proven at trial.'
    )

    path = os.path.join(OUTPUT_DIR, "brief_c_mixed_real_fake.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────
# BRIEF D: All Real Cases — Employment Discrimination (Title VII)
# Expected: All citations verified, characterizations accurate
# ─────────────────────────────────────────────────────────────────────────

def create_brief_d():
    doc = Document()

    _add_heading(doc, "MEMORANDUM IN SUPPORT OF MOTION FOR SUMMARY JUDGMENT")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Davis v. Consolidated Industries, Inc.")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    _add_heading(doc, "I. THE McDONNELL DOUGLAS BURDEN-SHIFTING FRAMEWORK", level=2)

    _add_para(doc,
        'Employment discrimination claims under Title VII of the Civil Rights Act '
        'of 1964 are evaluated using the burden-shifting framework established in '
        'McDonnell Douglas Corp. v. Green, 411 U.S. 792 (1973). Under this framework, '
        'the plaintiff must first establish a prima facie case of discrimination by '
        'showing membership in a protected class, qualification for the position, an '
        'adverse employment action, and circumstances giving rise to an inference of '
        'discrimination. Id. at 802.'
    )

    _add_para(doc,
        'Once the plaintiff establishes a prima facie case, the burden shifts to the '
        'employer to articulate a legitimate, nondiscriminatory reason for the adverse '
        'action. Texas Department of Community Affairs v. Burdine, 450 U.S. 248 (1981). '
        'Critically, the Court clarified that this is a burden of production, not '
        'persuasion:'
    )

    _add_block_quote(doc,
        '"The ultimate burden of persuading the trier of fact that the defendant '
        'intentionally discriminated against the plaintiff remains at all times '
        'with the plaintiff."'
    )

    _add_para(doc,
        'Texas Department of Community Affairs v. Burdine, 450 U.S. at 253.'
    )

    _add_heading(doc, "II. DISPARATE IMPACT DISCRIMINATION", level=2)

    _add_para(doc,
        'Title VII prohibits not only intentional discrimination but also facially '
        'neutral employment practices that have a discriminatory effect. Griggs v. '
        'Duke Power Co., 401 U.S. 424 (1971). The Court held:'
    )

    _add_block_quote(doc,
        '"The Act proscribes not only overt discrimination but also practices that '
        'are fair in form, but discriminatory in operation."'
    )

    _add_para(doc,
        'Griggs v. Duke Power Co., 401 U.S. at 431. The disparate impact framework '
        'was later codified in the Civil Rights Act of 1991 and further refined in '
        'Ricci v. DeStefano, 557 U.S. 557 (2009), where the Court held that before '
        'an employer can engage in intentional discrimination to avoid disparate impact '
        'liability, the employer must have a strong basis in evidence to believe it '
        'would be subject to disparate impact liability.'
    )

    _add_heading(doc, "III. PROTECTION AGAINST RETALIATION", level=2)

    _add_para(doc,
        'Title VII also prohibits retaliation against employees who oppose unlawful '
        'employment practices or participate in Title VII proceedings. Burlington '
        'Northern & Santa Fe Railway Co. v. White, 548 U.S. 53 (2006). The Court '
        'adopted a broad, objective standard for actionable retaliation, holding that '
        'a plaintiff must show that a reasonable employee would have found the '
        'challenged action materially adverse, "which in this context means it well '
        'might have dissuaded a reasonable worker from making or supporting a charge '
        'of discrimination." Id. at 68.'
    )

    _add_heading(doc, "IV. HOSTILE WORK ENVIRONMENT", level=2)

    _add_para(doc,
        'The Supreme Court recognized that workplace harassment creating a hostile '
        'or abusive working environment constitutes actionable employment '
        'discrimination under Title VII. Meritor Savings Bank, FSB v. Vinson, 477 '
        'U.S. 57 (1986). To establish a hostile work environment claim, the conduct '
        'must be sufficiently severe or pervasive to alter the conditions of '
        'employment and create an abusive working environment. Harris v. Forklift '
        'Systems, Inc., 510 U.S. 17 (1993). The Court explained that this '
        'determination requires examining the totality of the circumstances, '
        'including the frequency of the conduct, its severity, whether it is '
        'physically threatening or humiliating, and whether it unreasonably '
        'interferes with an employee\'s work performance. Id. at 23.'
    )

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'For the foregoing reasons, Plaintiff respectfully requests that this Court '
        'grant summary judgment on her Title VII claims.'
    )

    path = os.path.join(OUTPUT_DIR, "brief_d_employment_real.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────
# BRIEF E: Real Cases — Fabricated Quotes (Criminal Procedure / 4th Amend.)
# Expected: Citations found (cases are real), quotes/characterizations flagged
# ─────────────────────────────────────────────────────────────────────────

def create_brief_e():
    doc = Document()

    _add_heading(doc, "MEMORANDUM IN SUPPORT OF MOTION TO SUPPRESS EVIDENCE")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("State v. Williams")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    _add_heading(doc, "I. THE EXCLUSIONARY RULE", level=2)

    # Real case, FABRICATED quote — adds "automatically" and "regardless of
    # circumstances" which is wrong (there are exceptions like Leon good faith)
    _add_para(doc,
        'The exclusionary rule, as applied to the states through the Fourteenth '
        'Amendment, requires suppression of illegally obtained evidence. Mapp v. '
        'Ohio, 367 U.S. 643 (1961). The Court held that:'
    )

    _add_block_quote(doc,
        '"All evidence obtained by searches and seizures in violation of the '
        'Constitution is, by that same authority, automatically inadmissible in '
        'a state court regardless of the circumstances of the violation or the '
        'reliability of the evidence obtained."'
    )

    _add_para(doc, 'Mapp v. Ohio, 367 U.S. at 655.')

    _add_heading(doc, "II. REASONABLE EXPECTATION OF PRIVACY", level=2)

    # Real case, FABRICATED quote — drops the objective reasonableness prong
    # from Justice Harlan's concurrence (the actual test requires BOTH subjective
    # expectation AND society recognizing it as reasonable)
    _add_para(doc,
        'The Fourth Amendment protects people, not places. Katz v. United States, '
        '389 U.S. 347 (1967). Justice Harlan\'s concurrence, which became the '
        'governing standard, provides that:'
    )

    _add_block_quote(doc,
        '"Wherever a person has a subjective expectation of privacy, the Fourth '
        'Amendment provides absolute protection against government surveillance, '
        'regardless of whether that expectation is one that society would recognize '
        'as objectively reasonable."'
    )

    _add_para(doc, 'Katz v. United States, 389 U.S. at 361.')

    _add_heading(doc, "III. INVESTIGATORY STOPS AND SEARCHES", level=2)

    # Real case, FABRICATED characterization — Terry only allows a limited
    # pat-down of outer clothing for weapons, not "comprehensive searches"
    _add_para(doc,
        'When a police officer observes conduct that leads to a reasonable suspicion '
        'of criminal activity, the officer may conduct a comprehensive search of the '
        'person and their belongings. Terry v. Ohio, 392 U.S. 1 (1968). The Terry '
        'doctrine permits officers to search for any evidence of criminal activity '
        'during an investigatory stop, provided the officer can articulate specific '
        'facts justifying the intrusion.'
    )

    _add_heading(doc, "IV. PROBABLE CAUSE DETERMINATIONS", level=2)

    # Real case, FABRICATED characterization — Gates actually replaced the rigid
    # Aguilar-Spinelli test with a flexible totality-of-circumstances approach
    # that expressly allows hearsay and informant tips
    _add_para(doc,
        'The determination of probable cause requires that the affiant officer have '
        'direct, personal knowledge of the facts establishing criminal activity. '
        'Illinois v. Gates, 462 U.S. 213 (1983). The Court held that probable cause '
        'cannot be established through anonymous tips or hearsay alone, and that '
        'officers must independently verify each factual allegation before seeking '
        'a warrant. Id. at 238.'
    )

    _add_heading(doc, "V. DIGITAL PRIVACY", level=2)

    # Real case with CORRECT characterization (control case)
    _add_para(doc,
        'The Supreme Court has recognized that digital technology requires enhanced '
        'constitutional protections. Riley v. California, 573 U.S. 373 (2014). The '
        'Court held unanimously that police generally may not, without a warrant, '
        'search digital information on a cell phone seized from an individual who '
        'has been arrested. The Court reasoned that cell phones differ in both a '
        'quantitative and a qualitative sense from other objects that might be '
        'kept on an arrestee\'s person.'
    )

    _add_heading(doc, "VI. THE GOOD FAITH EXCEPTION", level=2)

    # Real case, FABRICATED quote — Leon requires OBJECTIVE good faith (reasonable
    # reliance on warrant from neutral magistrate), not SUBJECTIVE belief
    _add_para(doc,
        'Even where a warrant is later found deficient, evidence need not be '
        'suppressed under certain circumstances. United States v. Leon, 468 U.S. '
        '897 (1984). The Court stated:'
    )

    _add_block_quote(doc,
        '"The good faith exception applies whenever a law enforcement officer '
        'subjectively believes that the search being conducted is lawful, '
        'regardless of whether a reasonable officer would share that belief."'
    )

    _add_para(doc, 'United States v. Leon, 468 U.S. at 922.')

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'For the foregoing reasons, Defendant respectfully requests that this Court '
        'suppress all evidence obtained in violation of the Fourth Amendment.'
    )

    path = os.path.join(OUTPUT_DIR, "brief_e_criminal_fabricated.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────
# BRIEF F: Mix of Real and Hallucinated Cases — IP & Class Actions
# Expected: Real cases verified; fake cases flagged as errors/unverifiable
# ─────────────────────────────────────────────────────────────────────────

def create_brief_f():
    doc = Document()

    _add_heading(doc, "MEMORANDUM IN SUPPORT OF MOTION FOR CLASS CERTIFICATION")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Parker v. TechStream Solutions, Inc.")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    _add_heading(doc, "I. PATENT ELIGIBILITY UNDER SECTION 101", level=2)

    # REAL case — verifiable
    _add_para(doc,
        'The Supreme Court established a two-step framework for determining patent '
        'eligibility under 35 U.S.C. § 101. Alice Corp. v. CLS Bank International, '
        '573 U.S. 208 (2014). First, the court must determine whether the claims are '
        'directed to a patent-ineligible concept such as an abstract idea. Second, '
        'the court must examine the elements of the claim to determine whether they '
        'contain an inventive concept sufficient to transform the nature of the claim '
        'into a patent-eligible application.'
    )

    # HALLUCINATED case — does not exist
    _add_para(doc,
        'The Federal Circuit applied the Alice framework to software-implemented '
        'inventions in Sterling Technologies, Inc. v. DataVault Corp., 789 F.3d 432 '
        '(Fed. Cir. 2015), holding that "a software patent claim directed to the '
        'abstract idea of organizing and storing data in a hierarchical structure '
        'lacks the inventive concept necessary to survive scrutiny under Section 101." '
        'Id. at 441.'
    )

    _add_heading(doc, "II. FAIR USE IN COPYRIGHT", level=2)

    # REAL case — verifiable
    _add_para(doc,
        'The fair use doctrine permits certain uses of copyrighted material without '
        'authorization. Campbell v. Acuff-Rose Music, Inc., 510 U.S. 569 (1994). '
        'The Court held that 2 Live Crew\'s commercial parody of Roy Orbison\'s '
        '"Oh, Pretty Woman" could constitute fair use, and emphasized that '
        'transformative use is the central inquiry under the first statutory factor.'
    )

    # HALLUCINATED case — does not exist
    _add_para(doc,
        'The Sixth Circuit extended the transformative use doctrine to digital '
        'content aggregation in Henderson v. Creative Digital Solutions, LLC, 634 '
        'F.3d 891 (6th Cir. 2011). The court held that "automated compilation and '
        'indexing of copyrighted digital content constitutes transformative fair use '
        'when it serves a fundamentally different purpose than the original work." '
        'Id. at 899.'
    )

    _add_heading(doc, "III. INJUNCTIVE RELIEF IN PATENT CASES", level=2)

    # REAL case — verifiable
    _add_para(doc,
        'The Supreme Court rejected the Federal Circuit\'s categorical rule that '
        'permanent injunctions should issue upon a finding of patent infringement. '
        'eBay Inc. v. MercExchange, L.L.C., 547 U.S. 388 (2006). Instead, the '
        'Court held that the traditional four-factor test for equitable relief '
        'applies: (1) irreparable injury, (2) inadequacy of legal remedies, '
        '(3) balance of hardships, and (4) public interest. Id. at 391.'
    )

    # HALLUCINATED case — does not exist
    _add_para(doc,
        'The Ninth Circuit applied the eBay framework to preliminary injunctions '
        'in intellectual property disputes in Marcus v. Silicon Valley Innovations '
        'Corp., 856 F.3d 1124 (9th Cir. 2017), holding that "the traditional '
        'four-factor test governs preliminary injunctive relief in all intellectual '
        'property cases, whether patent, copyright, or trade secret." Id. at 1132.'
    )

    _add_heading(doc, "IV. CLASS CERTIFICATION REQUIREMENTS", level=2)

    # REAL case — verifiable
    _add_para(doc,
        'Class certification under Federal Rule of Civil Procedure 23 requires '
        'satisfaction of the prerequisites of numerosity, commonality, typicality, '
        'and adequacy. Wal-Mart Stores, Inc. v. Dukes, 564 U.S. 338 (2011). The '
        'Court tightened the commonality requirement, holding that plaintiffs must '
        'demonstrate that the class members have suffered the same injury, and '
        'that their claims depend upon a common contention that is capable of '
        'classwide resolution.'
    )

    # HALLUCINATED case — does not exist
    _add_para(doc,
        'The Third Circuit applied Dukes to consumer class actions in Porter v. '
        'National Financial Services Corp., 912 F.3d 267 (3d Cir. 2019). The '
        'court held that "where a defendant\'s uniform policy affects all class '
        'members identically, commonality is satisfied even if individual damages '
        'vary." Id. at 278.'
    )

    _add_heading(doc, "V. ARBITRATION AND CLASS WAIVERS", level=2)

    # REAL case with MISLEADING characterization — AT&T Mobility held that the
    # FAA preempts state laws that condition enforcement of arbitration on
    # availability of class procedures; it does NOT broadly "prohibit all class
    # actions" as stated here
    _add_para(doc,
        'The Federal Arbitration Act broadly prohibits all class action proceedings '
        'when a valid arbitration agreement exists between the parties. AT&T '
        'Mobility LLC v. Concepcion, 563 U.S. 333 (2011). The Court held that '
        'the FAA mandates individual arbitration in all consumer disputes, '
        'effectively eliminating the class action mechanism whenever an arbitration '
        'clause is present.'
    )

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'For the foregoing reasons, Plaintiff respectfully requests that this Court '
        'certify the proposed class and enjoin Defendant from further infringement '
        'of Plaintiff\'s intellectual property rights.'
    )

    path = os.path.join(OUTPUT_DIR, "brief_f_ip_mixed.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────
# BRIEF G: Stress Test — Constitutional Law / Civil Rights
# 18 citations: 10 real, 5 fake, 3 real with fabricated content
# Tests: landmark SCOTUS cases, abbreviation handling, dense citation load
# ─────────────────────────────────────────────────────────────────────────

def create_brief_g():
    doc = Document()

    _add_heading(doc, "MEMORANDUM IN SUPPORT OF MOTION FOR PRELIMINARY INJUNCTION")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Anderson v. State of Columbia Board of Education")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    _add_heading(doc, "I. JUDICIAL REVIEW AND CONSTITUTIONAL FRAMEWORK", level=2)

    # REAL — Marbury v. Madison
    _add_para(doc,
        'The power of judicial review is foundational to our constitutional system. '
        'As Chief Justice Marshall declared in Marbury v. Madison, 5 U.S. 137 (1803):'
    )

    _add_block_quote(doc,
        '"It is emphatically the province and duty of the judicial department to say '
        'what the law is."'
    )

    _add_para(doc,
        'Marbury v. Madison, 5 U.S. at 177. This principle has guided constitutional '
        'adjudication for over two centuries.'
    )

    # FAKE — Patterson v. State Board of Education
    _add_para(doc,
        'The scope of judicial review extends to all state educational policies that '
        'implicate fundamental rights. Patterson v. State Bd. of Educ., 543 F.3d 891 '
        '(5th Cir. 2008) (holding that "federal courts have an independent obligation '
        'to scrutinize state educational policies that burden constitutional rights, '
        'regardless of the state\'s asserted pedagogical justification"). Id. at 903.'
    )

    _add_heading(doc, "II. EQUAL PROTECTION", level=2)

    # REAL — Brown v. Board of Education
    _add_para(doc,
        'The Equal Protection Clause of the Fourteenth Amendment prohibits state-sponsored '
        'racial segregation. Brown v. Board of Education, 347 U.S. 483 (1954). The Court '
        'unanimously held that "in the field of public education, the doctrine of \'separate '
        'but equal\' has no place. Separate educational facilities are inherently unequal." '
        'Id. at 495.'
    )

    # REAL — Loving v. Virginia
    _add_para(doc,
        'Racial classifications are subject to strict scrutiny and must be justified by '
        'a compelling governmental interest. Loving v. Virginia, 388 U.S. 1 (1967). The '
        'Court struck down anti-miscegenation statutes, holding that the freedom to marry '
        'is a vital personal right essential to the orderly pursuit of happiness by free '
        'men. Id. at 12.'
    )

    # REAL — Obergefell v. Hodges
    _add_para(doc,
        'The fundamental right to marry extends to same-sex couples under both the Due '
        'Process and Equal Protection Clauses. Obergefell v. Hodges, 576 U.S. 644 (2015). '
        'The Court identified four principles and traditions demonstrating that the reasons '
        'marriage is fundamental apply with equal force to same-sex couples.'
    )

    # FAKE — Davidson v. Metro. Housing Auth.
    _add_para(doc,
        'Equal protection principles apply with particular force in the housing context. '
        'Davidson v. Metro. Housing Auth., 723 F.3d 556 (6th Cir. 2013). The Sixth '
        'Circuit held that "municipal housing authorities may not adopt facially neutral '
        'policies whose primary purpose and predictable effect is to perpetuate racial '
        'segregation in public housing." Id. at 567.'
    )

    _add_heading(doc, "III. FIRST AMENDMENT PROTECTIONS", level=2)

    # REAL — New York Times v. Sullivan
    _add_para(doc,
        'The First Amendment imposes stringent limits on defamation claims brought by '
        'public officials. New York Times Co. v. Sullivan, 376 U.S. 254 (1964). The '
        'Court held that a public official must prove that the defamatory statement was '
        'made with "actual malice" — that is, with knowledge that it was false or with '
        'reckless disregard of whether it was false or not. Id. at 279-80.'
    )

    # REAL — Tinker v. Des Moines
    _add_para(doc,
        'Students retain their constitutional rights within the school setting. Tinker '
        'v. Des Moines Indep. Cmty. Sch. Dist., 393 U.S. 503 (1969). The Court '
        'famously declared that students do not "shed their constitutional rights to '
        'freedom of speech or expression at the schoolhouse gate." Id. at 506.'
    )

    # REAL — Brandenburg v. Ohio
    _add_para(doc,
        'The government may not prohibit advocacy of illegal conduct unless such advocacy '
        'is "directed to inciting or producing imminent lawless action and is likely to '
        'incite or produce such action." Brandenburg v. Ohio, 395 U.S. 444, 447 (1969).'
    )

    # REAL — Citizens United v. FEC
    _add_para(doc,
        'Political speech does not lose First Amendment protection simply because its '
        'source is a corporation. Citizens United v. Federal Election Comm\'n, 558 U.S. '
        '310 (2010). The Court held that the First Amendment prohibits Congress from '
        'restricting independent political expenditures by corporations and other '
        'associations.'
    )

    # FAKE — Whitmore v. City of Portland
    _add_para(doc,
        'Public forum protections extend to digital spaces maintained by government '
        'entities. Whitmore v. City of Portland, 678 F.3d 1205 (9th Cir. 2012) (holding '
        'that "a government-operated social media page constitutes a designated public '
        'forum subject to traditional First Amendment constraints"). Id. at 1214.'
    )

    # FAKE — Reynolds v. Federal Election Comm'n
    _add_para(doc,
        'The D.C. Circuit has also recognized that disclosure requirements imposed on '
        'political organizations may violate the First Amendment when they chill the '
        'exercise of associational freedoms. Reynolds v. Fed. Election Comm\'n, 891 F.3d '
        '445 (D.C. Cir. 2018).'
    )

    # REAL/FABRICATED — Reed v. Town of Gilbert (altered — says "all sign regulation"
    # violates 1A, but Reed actually held only content-based restrictions get strict scrutiny)
    _add_para(doc,
        'All government regulation of signage constitutes a per se violation of the '
        'First Amendment. Reed v. Town of Gilbert, 576 U.S. 155 (2015). The Court held '
        'that municipalities may not impose any restrictions on the content, size, or '
        'placement of signs without violating the free speech guarantee.'
    )

    _add_heading(doc, "IV. FUNDAMENTAL RIGHTS AND INCORPORATION", level=2)

    # REAL — Gideon v. Wainwright
    _add_para(doc,
        'The Sixth Amendment right to counsel is a fundamental right applicable to the '
        'states through the Fourteenth Amendment. Gideon v. Wainwright, 372 U.S. 335 '
        '(1963). The Court held that any person charged with a serious criminal offense '
        'who cannot afford an attorney must be provided one at government expense.'
    )

    # REAL — District of Columbia v. Heller
    _add_para(doc,
        'The Second Amendment protects an individual right to possess firearms '
        'unconnected with service in a militia, and to use that arm for traditionally '
        'lawful purposes such as self-defense within the home. District of Columbia v. '
        'Heller, 554 U.S. 570 (2008).'
    )

    # REAL/FABRICATED — Shelby County v. Holder (fabricated quote —
    # real holding struck down coverage formula, not all preclearance)
    _add_para(doc,
        'The Voting Rights Act\'s preclearance requirement was struck down in its '
        'entirety as unconstitutional. Shelby County v. Holder, 570 U.S. 529 (2013). '
        'The Court declared:'
    )

    _add_block_quote(doc,
        '"All preclearance requirements imposed under the Voting Rights Act are '
        'permanently unconstitutional because they violate the fundamental principle '
        'of equal sovereignty among the states and impose burdens that bear no '
        'rational relationship to current conditions."'
    )

    _add_para(doc, 'Shelby County v. Holder, 570 U.S. at 556.')

    # REAL/FABRICATED — McDonald v. City of Chicago (fabricated characterization —
    # real holding was incorporation, not public carry rights)
    _add_para(doc,
        'The Second Amendment right extends beyond the home and establishes an absolute '
        'right to carry firearms in public spaces. McDonald v. City of Chicago, 561 U.S. '
        '742 (2010). The Court held that states and municipalities have no authority to '
        'restrict the open carrying of firearms in any public area.'
    )

    # FAKE — Crawford v. State Dep't of Educ.
    _add_para(doc,
        'See also Crawford v. State Dep\'t of Educ., 456 F.3d 789 (8th Cir. 2006) '
        '(holding that "the fundamental right to direct one\'s children\'s education '
        'requires heightened scrutiny of any state regulation that impinges upon '
        'parental choice in schooling").'
    )

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'For the foregoing reasons, Plaintiff respectfully requests that this Court '
        'grant the motion for a preliminary injunction.'
    )

    path = os.path.join(OUTPUT_DIR, "brief_g_constitutional_stress.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────
# BRIEF H: Stress Test — Commercial Litigation / Securities Fraud
# 16 citations: 8 real, 5 fake, 3 real with fabricated content
# Tests: abbreviation-heavy names (R.R., Pharms., Assocs., Ins.), dense cites
# ─────────────────────────────────────────────────────────────────────────

def create_brief_h():
    doc = Document()

    _add_heading(doc, "MEMORANDUM IN SUPPORT OF MOTION FOR CLASS CERTIFICATION")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Wellington Capital Fund, LP v. Pacific Tech Holdings, Inc.")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    _add_heading(doc, "I. PERSONAL JURISDICTION AND CHOICE OF LAW", level=2)

    # REAL — International Shoe (abbreviation test: Co.)
    _add_para(doc,
        'The exercise of personal jurisdiction must comport with due process. '
        'International Shoe Co. v. Washington, 326 U.S. 310 (1945). The Court held '
        'that a defendant must have "minimum contacts" with the forum state "such that '
        'the maintenance of the suit does not offend traditional notions of fair play '
        'and substantial justice." Id. at 316.'
    )

    # REAL — Erie R.R. (abbreviation test: R.R., Co.)
    _add_para(doc,
        'In diversity cases, federal courts must apply state substantive law. Erie R.R. '
        'Co. v. Tompkins, 304 U.S. 64 (1938). The Erie doctrine ensures that the '
        'outcome of litigation does not depend on whether the case is adjudicated in '
        'state or federal court.'
    )

    # REAL — Shady Grove (abbreviation test: Assocs., Ins., Co.)
    _add_para(doc,
        'Federal Rule of Civil Procedure 23 governs class certification in federal '
        'court, even when state law would prohibit class actions for the claim at issue. '
        'Shady Grove Orthopedic Assocs., P.A. v. Allstate Ins. Co., 559 U.S. 393 (2010).'
    )

    # FAKE — Thornton v. Pacific Coast Fin. Group (abbreviation-laden fake)
    _add_para(doc,
        'The Ninth Circuit has applied these principles in the securities context, '
        'holding that minimum contacts are established when a defendant directs '
        'fraudulent communications into the forum. Thornton v. Pacific Coast Fin. '
        'Group, 823 F.3d 1156 (9th Cir. 2016).'
    )

    # FAKE — Westbrook Inv. Corp. v. Atlantic Mgmt. Servs.
    _add_para(doc,
        'See also Westbrook Inv. Corp. v. Atlantic Mgmt. Servs., 567 F. Supp. 2d 432 '
        '(S.D.N.Y. 2008) (applying Erie to determine that New York substantive law '
        'governs claims arising from investment management agreements executed in '
        'New York).'
    )

    _add_heading(doc, "II. SECURITIES FRAUD UNDER SECTION 10(b)", level=2)

    # REAL — Basic Inc. v. Levinson
    _add_para(doc,
        'The fraud-on-the-market theory creates a rebuttable presumption of reliance '
        'for plaintiffs alleging securities fraud. Basic Inc. v. Levinson, 485 U.S. 224 '
        '(1988). The Court held that "the market price of shares traded on well-developed '
        'markets reflects all publicly available information, and, hence, any material '
        'misrepresentations." Id. at 246.'
    )

    # REAL — Tellabs (abbreviation test: Inc., Ltd.)
    _add_para(doc,
        'A securities fraud complaint must state facts giving rise to a "strong inference" '
        'of scienter — an inference "at least as compelling as any opposing inference one '
        'could draw from the facts alleged." Tellabs, Inc. v. Makor Issues & Rights, '
        'Ltd., 551 U.S. 308, 324 (2007).'
    )

    # REAL — Stoneridge Inv. Partners
    _add_para(doc,
        'Section 10(b) does not create a private right of action against secondary actors '
        'who did not make a public misstatement or omission upon which investors relied. '
        'Stoneridge Inv. Partners, LLC v. Scientific-Atlanta, Inc., 552 U.S. 148 (2008).'
    )

    # REAL — Dura Pharmaceuticals (abbreviation test: Pharms., Inc.)
    _add_para(doc,
        'Plaintiffs must demonstrate loss causation — that the defendant\'s fraud actually '
        'caused an economic loss, not merely an inflated purchase price. Dura Pharms., '
        'Inc. v. Broudo, 544 U.S. 336 (2005).'
    )

    # REAL/FABRICATED — Halliburton (fabricated: says it "abolished" the fraud-on-market
    # presumption; actually it reaffirmed the presumption but allowed rebuttal at cert.)
    _add_para(doc,
        'The Supreme Court abolished the fraud-on-the-market presumption of reliance, '
        'holding that plaintiffs must demonstrate actual, individualized reliance on '
        'specific misstatements. Halliburton Co. v. Erica P. John Fund, Inc., 573 U.S. '
        '258 (2014). The Court stated that "the fraud-on-the-market theory has no place '
        'in modern securities litigation and is hereby overruled." Id. at 277.'
    )

    # REAL/FABRICATED — Morrison v. National Australia Bank (fabricated: says 10(b) has
    # no territorial limits; actually held it applies ONLY domestically)
    _add_para(doc,
        'Section 10(b) of the Exchange Act has no territorial limitations and applies '
        'to securities transactions worldwide, including those occurring entirely on '
        'foreign exchanges. Morrison v. Nat\'l Australia Bank Ltd., 561 U.S. 247 (2010).'
    )

    # FAKE — Meridian Capital Partners
    _add_para(doc,
        'See Meridian Capital Partners v. Blackstone Advisors, LLC, 745 F.3d 328 (2d '
        'Cir. 2014) (holding that "where a defendant\'s fraudulent scheme spans multiple '
        'jurisdictions, loss causation may be established through aggregate market impact '
        'analysis rather than transaction-specific proof").'
    )

    # FAKE — Sullivan v. Global Tech. Innovations
    _add_para(doc,
        'The Third Circuit has similarly held that scienter may be inferred from '
        'circumstantial evidence of a corporate officer\'s access to contradictory '
        'information. Sullivan v. Global Tech. Innovations, Inc., 901 F.3d 178 (3d '
        'Cir. 2018).'
    )

    _add_heading(doc, "III. CLASS CERTIFICATION AND STANDING", level=2)

    # REAL — Blue Chip Stamps
    _add_para(doc,
        'Only actual purchasers or sellers of securities have standing to bring a '
        'private action under Rule 10b-5. Blue Chip Stamps v. Manor Drug Stores, '
        '421 U.S. 723 (1975). This purchaser-seller requirement serves to prevent '
        'vexatious litigation by those who had no investment at stake.'
    )

    # REAL/FABRICATED — Janus Capital (fabricated: says anyone who "contributes" is liable;
    # actually narrowed maker liability to those with "ultimate authority" over the statement)
    _add_para(doc,
        'Any person who contributes to the creation of a fraudulent misstatement may '
        'be held liable as the "maker" of that statement under Rule 10b-5. Janus Capital '
        'Group, Inc. v. First Derivative Traders, 564 U.S. 135 (2011). The Court broadly '
        'defined "maker" to include anyone who participates in drafting, editing, or '
        'approving the misleading statement.'
    )

    # FAKE — Chen v. Pacific Semiconductor Corp.
    _add_para(doc,
        'See also Chen v. Pacific Semiconductor Corp., 678 F.3d 923 (9th Cir. 2012) '
        '(holding that class certification is appropriate where common questions of '
        'scienter and materiality predominate over individual questions of reliance '
        'in a securities fraud class action).'
    )

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'For the foregoing reasons, Plaintiff respectfully requests that this Court '
        'certify the proposed class under Federal Rule of Civil Procedure 23(b)(3).'
    )

    path = os.path.join(OUTPUT_DIR, "brief_h_securities_stress.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────
# BRIEF I: Stress Test — Administrative Law / Federal Procedure
# 16 citations: 8 real, 5 fake, 3 real with fabricated content
# Tests: heavily abbreviated names (Mfrs., Ass'n, Mut., Pharm., Envtl.)
# ─────────────────────────────────────────────────────────────────────────

def create_brief_i():
    doc = Document()

    _add_heading(doc, "MEMORANDUM IN SUPPORT OF MOTION FOR SUMMARY JUDGMENT")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("National Environmental Coalition v. Environmental Protection Agency")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    _add_heading(doc, "I. ARTICLE III STANDING", level=2)

    # REAL — Lujan v. Defenders of Wildlife
    _add_para(doc,
        'Article III standing requires three elements: (1) an injury in fact that is '
        'concrete and particularized, (2) a causal connection between the injury and the '
        'challenged conduct, and (3) a likelihood that a favorable judicial decision will '
        'redress the injury. Lujan v. Defenders of Wildlife, 504 U.S. 555, 560-61 (1992).'
    )

    # REAL — Spokeo v. Robins
    _add_para(doc,
        'The injury must be "concrete" — a bare procedural violation of a statute, '
        'without any concrete harm, does not satisfy Article III. Spokeo, Inc. v. Robins, '
        '578 U.S. 330 (2016).'
    )

    # FAKE — Cartwright v. Dep't of Health & Human Servs.
    _add_para(doc,
        'Environmental organizations have standing to challenge agency action when their '
        'members use the affected natural resources. Cartwright v. Dep\'t of Health & '
        'Human Servs., 634 F.3d 782 (4th Cir. 2011) (holding that "organizational '
        'standing exists where an agency\'s regulatory failure directly impairs the '
        'organization\'s mission and diverts its resources").'
    )

    # REAL/FABRICATED — Massachusetts v. EPA (fabricated: says private citizens have
    # same standing as states; actually only states had quasi-sovereign standing)
    _add_para(doc,
        'Private citizens have the same standing as states to challenge the EPA\'s '
        'failure to regulate greenhouse gas emissions under the Clean Air Act. '
        'Massachusetts v. EPA, 549 U.S. 497 (2007). The Court held that any party '
        'demonstrating exposure to the effects of climate change satisfies the '
        'injury-in-fact requirement without any special showing.'
    )

    _add_heading(doc, "II. STANDARD OF REVIEW FOR AGENCY ACTION", level=2)

    # REAL — Chevron (abbreviation test: U.S.A., Inc.)
    _add_para(doc,
        'Courts review agency interpretations of ambiguous statutes under the two-step '
        'framework established in Chevron U.S.A., Inc. v. Natural Resources Defense '
        'Council, Inc., 467 U.S. 837 (1984). First, the court asks whether Congress '
        'has directly spoken to the precise question at issue. If the statute is '
        'ambiguous, the court defers to the agency\'s interpretation so long as it is '
        'a permissible construction. Id. at 842-43.'
    )

    # REAL — Auer v. Robbins
    _add_para(doc,
        'An agency\'s interpretation of its own ambiguous regulation is controlling '
        'unless plainly erroneous or inconsistent with the regulation. Auer v. Robbins, '
        '519 U.S. 452, 461 (1997).'
    )

    # REAL — Motor Vehicle Mfrs. (abbreviation stress test: Mfrs., Ass'n, Mut., Auto., Ins., Co.)
    _add_para(doc,
        'Agency action is "arbitrary and capricious" if the agency relied on factors '
        'Congress did not intend, entirely failed to consider an important aspect of the '
        'problem, offered an explanation that runs counter to the evidence before it, or '
        'the action is so implausible that it could not be ascribed to a difference in '
        'view. Motor Vehicle Mfrs. Ass\'n of the U.S., Inc. v. State Farm Mut. Auto. '
        'Ins. Co., 463 U.S. 29, 43 (1983).'
    )

    # FAKE — Lexington Envtl. Servs. v. EPA (abbreviation-laden fake)
    _add_para(doc,
        'The Tenth Circuit has applied State Farm\'s "hard look" standard to EPA '
        'rulemaking, holding that the agency must provide a detailed explanation for '
        'any departure from prior practice. Lexington Envtl. Servs. v. EPA, 789 F.3d '
        '1034 (10th Cir. 2015).'
    )

    # REAL/FABRICATED — Kisor v. Wilkie (fabricated: says Court overruled Auer;
    # actually narrowed but preserved it)
    _add_para(doc,
        'The Supreme Court overruled Auer v. Robbins, holding that courts may never '
        'defer to an agency\'s interpretation of its own regulations. Kisor v. Wilkie, '
        '588 U.S. 558 (2019). The Court held that all agency regulatory interpretations '
        'must be reviewed de novo without any deference.'
    )

    # FAKE — Nakamura v. Fed. Trade Comm'n (abbreviation test)
    _add_para(doc,
        'See Nakamura v. Fed. Trade Comm\'n, 845 F.3d 267 (7th Cir. 2017) (holding '
        'that "an agency\'s failure to provide adequate notice of a proposed rule change '
        'renders the final rule void ab initio under the Administrative Procedure Act").'
    )

    _add_heading(doc, "III. AGENCY AUTHORITY AND PROCEDURE", level=2)

    # REAL/FABRICATED — FCC v. Fox Television (fabricated quote about "unfettered
    # discretion"; actually held agencies MUST provide reasoned explanation for changes)
    _add_para(doc,
        'Agencies have unfettered discretion to reverse prior policies without providing '
        'any explanation for the change. FCC v. Fox Television Stations, Inc., 556 U.S. '
        '502 (2009). The Court stated:'
    )

    _add_block_quote(doc,
        '"An agency need not demonstrate that the reasons for its new policy are '
        'better than the reasons for its old policy, and indeed need not acknowledge '
        'the change at all, so long as the new policy falls within the scope of the '
        'agency\'s delegated authority."'
    )

    _add_para(doc, 'FCC v. Fox Television Stations, Inc., 556 U.S. at 515.')

    # FAKE — Westfield Pharm. Corp. v. FDA (abbreviation test: Pharm., Corp.)
    _add_para(doc,
        'See Westfield Pharm. Corp. v. FDA, 912 F.3d 456 (D.C. Cir. 2019) (holding '
        'that "the FDA\'s expedited approval pathway does not exempt the agency from '
        'the notice-and-comment requirements of the APA when imposing post-market '
        'surveillance obligations").'
    )

    # FAKE — O'Brien v. Nat'l Labor Relations Bd.
    _add_para(doc,
        'Exhaustion of administrative remedies is jurisdictional and cannot be waived. '
        'O\'Brien v. Nat\'l Labor Relations Bd., 778 F.3d 1124 (11th Cir. 2015).'
    )

    _add_heading(doc, "IV. PERSONAL JURISDICTION OVER FEDERAL AGENCIES", level=2)

    # REAL — Bristol-Myers Squibb (abbreviation test: Co.)
    _add_para(doc,
        'Specific personal jurisdiction requires an affiliation between the forum and '
        'the underlying controversy. Bristol-Myers Squibb Co. v. Superior Court of '
        'California, 582 U.S. 255 (2017). The Court held that a defendant\'s general '
        'connections with the forum are insufficient to support specific jurisdiction '
        'over claims unrelated to those connections.'
    )

    # REAL — Daimler AG v. Bauman
    _add_para(doc,
        'General personal jurisdiction exists only where the defendant is "at home" — '
        'typically the state of incorporation or principal place of business. Daimler AG '
        'v. Bauman, 571 U.S. 117 (2014). The Court rejected the notion that a '
        'corporation\'s substantial business in a forum renders it subject to general '
        'jurisdiction there.'
    )

    # REAL — Piper Aircraft (abbreviation test: Co.)
    _add_para(doc,
        'Under the doctrine of forum non conveniens, courts balance private and public '
        'interest factors in deciding whether to dismiss in favor of a more convenient '
        'forum. Piper Aircraft Co. v. Reyno, 454 U.S. 235 (1981). The Court held that '
        'the possibility of an unfavorable change in law should ordinarily not be given '
        'conclusive or even substantial weight in the analysis.'
    )

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'For the foregoing reasons, Plaintiff respectfully requests that this Court '
        'grant summary judgment and vacate the EPA\'s final rule as arbitrary and '
        'capricious.'
    )

    path = os.path.join(OUTPUT_DIR, "brief_i_admin_stress.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────
# BRIEF J: EDGE CASES — Citation Format Challenges
# Tests: per curiam, "In re" cases, "et al.", string citations,
#   signal prefixes (See, Cf., But see), very old / very new cases,
#   parenthetical explanations, dissent cited correctly as dissent
# ─────────────────────────────────────────────────────────────────────────

def create_brief_j():
    doc = Document()

    _add_heading(doc, "MEMORANDUM IN OPPOSITION TO MOTION TO DISMISS")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("In re: Global Logistics Corp. Shareholder Litigation")
    run.bold = True

    _add_heading(doc, "I. PER CURIAM AND UNNAMED PARTY DECISIONS", level=2)

    # REAL — Bush v. Gore (per curiam, no named author)
    _add_para(doc,
        'The principle that standardless manual recounts violate equal protection '
        'was established by a per curiam decision of the full Court. Bush v. Gore, '
        '531 U.S. 98 (2000) (per curiam). The Court held that "the recount '
        'mechanisms implemented in response to the decisions of the Florida Supreme '
        'Court do not satisfy the minimum requirement for non-arbitrary treatment '
        'of voters necessary to secure the fundamental right."'
    )

    # REAL — In re Winship (In re case format)
    _add_para(doc,
        'The "beyond a reasonable doubt" standard for criminal cases was '
        'constitutionalized in In re Winship, 397 U.S. 358, 364 (1970), where '
        'the Court held that "the Due Process Clause protects the accused against '
        'conviction except upon proof beyond a reasonable doubt of every fact '
        'necessary to constitute the crime with which he is charged."'
    )

    # REAL — Roe v. Wade (overruled — test whether system notes this)
    _add_para(doc,
        'The right to privacy encompasses a woman\'s decision whether to terminate '
        'a pregnancy. Roe v. Wade, 410 U.S. 113 (1973). The Court established the '
        'trimester framework, holding that the State\'s interest becomes compelling '
        'at the point of viability.'
    )

    _add_heading(doc, "II. STRING CITATIONS AND SIGNAL PREFIXES", level=2)

    # REAL — String citation: three real cases with signal prefixes
    _add_para(doc,
        'Courts have long recognized the importance of procedural due process '
        'before deprivation of life, liberty, or property. See Mathews v. '
        'Eldridge, 424 U.S. 319, 335 (1976) (establishing three-factor balancing '
        'test); see also Cleveland Bd. of Educ. v. Loudermill, 470 U.S. 532 '
        '(1985) (holding that a public employee with a property interest in '
        'continued employment is entitled to a pre-termination hearing); cf. '
        'Goldberg v. Kelly, 397 U.S. 254, 264 (1970) (requiring an evidentiary '
        'hearing before termination of welfare benefits).'
    )

    # FAKE — Inserted into a string citation to test detection
    _add_para(doc,
        'The doctrine of qualified immunity shields government officials from '
        'civil damages liability unless the official violated a statutory or '
        'constitutional right that was clearly established at the time. See '
        'Harlow v. Fitzgerald, 457 U.S. 800, 818 (1982); Anderson v. Creighton, '
        '483 U.S. 635 (1987); Whitfield v. Mun. Auth. of Camden, 734 F.3d 291 '
        '(3d Cir. 2013) (holding that the "clearly established" prong requires '
        '"a robust consensus of cases of persuasive authority").'
    )

    _add_heading(doc, "III. DISSENT CITED AS DISSENT", level=2)

    # REAL — Korematsu dissent (Jackson, J., dissenting) — correctly attributed
    _add_para(doc,
        'As Justice Jackson warned in his celebrated dissent, a judicial validation '
        'of racial discrimination "lies about like a loaded weapon, ready for the '
        'hand of any authority that can bring forward a plausible claim of an '
        'urgent need." Korematsu v. United States, 323 U.S. 214, 246 (1944) '
        '(Jackson, J., dissenting).'
    )

    # REAL — Plessy v. Ferguson (Harlan, J., dissenting)
    _add_para(doc,
        'Justice Harlan\'s lone dissent foresaw the moral bankruptcy of the '
        'separate-but-equal doctrine: "Our Constitution is color-blind, and '
        'neither knows nor tolerates classes among citizens." Plessy v. Ferguson, '
        '163 U.S. 537, 559 (1896) (Harlan, J., dissenting).'
    )

    _add_heading(doc, "IV. PARENTHETICAL-HEAVY CITATIONS", level=2)

    # REAL — Daubert (parenthetical with full explanation)
    _add_para(doc,
        'Expert testimony must meet threshold reliability requirements. Daubert '
        'v. Merrell Dow Pharms., Inc., 509 U.S. 579 (1993) (replacing the '
        'Frye "general acceptance" test with a multi-factor reliability inquiry '
        'under Federal Rule of Evidence 702). The trial judge serves as a '
        '"gatekeeper" ensuring that expert testimony rests on a reliable '
        'foundation and is relevant to the task at hand. Id. at 597.'
    )

    # FAKE — Realistic-looking case embedded in parenthetical chain
    _add_para(doc,
        'District courts have broad discretion in managing discovery. See Hickman '
        'v. Taylor, 329 U.S. 495 (1947) (establishing work-product privilege); '
        'Oppenheimer Fund, Inc. v. Sanders, 437 U.S. 340 (1978) (discussing '
        'scope of permissible discovery under Rule 26); Gallagher v. Consolidated '
        'Freightways Corp. of Del., 789 F.3d 803 (8th Cir. 2015) (applying '
        'proportionality principles to limit burdensome ESI requests).'
    )

    _add_heading(doc, "V. VERY OLD AND RECENT CASES", level=2)

    # REAL — Marbury (1803, very old)
    _add_para(doc,
        'The judiciary has exercised the power of judicial review since the '
        'Republic\'s earliest days. Marbury v. Madison, 5 U.S. (1 Cranch) 137 '
        '(1803). Chief Justice Marshall declared that "[i]t is emphatically the '
        'province and duty of the judicial department to say what the law is."'
    )

    # REAL — Students for Fair Admissions (2023, very recent)
    _add_para(doc,
        'Race-conscious admissions programs at colleges and universities violate '
        'the Equal Protection Clause. Students for Fair Admissions, Inc. v. '
        'President & Fellows of Harvard Coll., 600 U.S. 181 (2023). The Court '
        'held that Harvard\'s and UNC\'s admissions programs lacked sufficiently '
        'focused and measurable objectives warranting the use of race.'
    )

    # REAL — Loper Bright (2024, very recent, overruling Chevron)
    _add_para(doc,
        'Courts must exercise their own judgment in deciding whether an agency '
        'has acted within its statutory authority. Loper Bright Enterprises v. '
        'Raimondo, 144 S. Ct. 2244 (2024). The Court overruled Chevron, holding '
        'that the Administrative Procedure Act requires courts to decide legal '
        'questions by applying their own judgment rather than deferring to '
        'agency interpretations.'
    )

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'For the foregoing reasons, Defendant\'s motion to dismiss should be '
        'denied in its entirety.'
    )

    path = os.path.join(OUTPUT_DIR, "brief_j_edge_cases.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────
# BRIEF K: SUBTLE FABRICATION TRAPS
# Tests: wrong volume number (off by one), dissent presented as majority,
#   holding of Case A attributed to Case B, real case wrong year,
#   overruled holding cited as current law, accurate quote but wrong case
# ─────────────────────────────────────────────────────────────────────────

def create_brief_k():
    doc = Document()

    _add_heading(doc, "PLAINTIFF'S TRIAL BRIEF")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Rivera v. National Healthcare Systems, Inc.")
    run.bold = True

    _add_heading(doc, "I. STANDARD OF REVIEW", level=2)

    # REAL (correct) — Celotex as control
    _add_para(doc,
        'Summary judgment is appropriate when there is no genuine dispute as to '
        'any material fact and the movant is entitled to judgment as a matter of '
        'law. Celotex Corp. v. Catrett, 477 U.S. 317, 322 (1986).'
    )

    # TRAP: Wrong volume — Heller is at 554 U.S. 570, NOT 555 U.S. 570
    _add_para(doc,
        'The Second Amendment protects an individual\'s right to possess a firearm '
        'unconnected with service in a militia and to use that arm for traditionally '
        'lawful purposes, such as self-defense within the home. District of Columbia '
        'v. Heller, 555 U.S. 570 (2008).'
    )

    _add_heading(doc, "II. DISSENT PRESENTED AS MAJORITY HOLDING", level=2)

    # TRAP: Shelby County — the dissent's position presented as if it were the majority
    _add_para(doc,
        'The Voting Rights Act\'s preclearance requirement remains a valid exercise '
        'of Congressional enforcement power under the Fifteenth Amendment. The '
        'coverage formula in Section 4(b) is justified by current conditions of '
        'voting discrimination. Shelby County v. Holder, 570 U.S. 529, 559 '
        '(2013). The Court emphasized that Congress had amassed a substantial '
        'record of ongoing discrimination sufficient to warrant the extraordinary '
        'remedy of preclearance.'
    )

    # TRAP: Citizens United — dissent's position presented as holding
    _add_para(doc,
        'Corporations do not possess the same First Amendment rights as natural '
        'persons, and the government may restrict corporate expenditures in '
        'elections to prevent corruption and the appearance of corruption. '
        'Citizens United v. Federal Election Comm\'n, 558 U.S. 310 (2010). '
        'The Court held that the government\'s interest in preventing the '
        '"distortion" of political debate by large corporate spending justified '
        'expenditure restrictions.'
    )

    _add_heading(doc, "III. HOLDING SWAPPED BETWEEN CASES", level=2)

    # TRAP: Miranda holding attributed to Terry v. Ohio
    _add_para(doc,
        'Before a custodial interrogation, officers must inform the suspect of '
        'their right to remain silent and their right to an attorney. Terry v. '
        'Ohio, 392 U.S. 1 (1968). This procedural safeguard is required to '
        'protect the Fifth Amendment privilege against self-incrimination.'
    )

    # TRAP: Terry's holding attributed to Miranda
    _add_para(doc,
        'An officer may conduct a brief investigatory stop when the officer has '
        'reasonable suspicion that criminal activity may be afoot, and may frisk '
        'the suspect for weapons if the officer reasonably believes the person is '
        'armed and dangerous. Miranda v. Arizona, 384 U.S. 436 (1966).'
    )

    _add_heading(doc, "IV. REAL QUOTE ATTRIBUTED TO WRONG CASE", level=2)

    # TRAP: Famous Marbury quote attributed to McCulloch v. Maryland
    _add_para(doc,
        'As the Court famously declared, "It is emphatically the province and '
        'duty of the judicial department to say what the law is." McCulloch v. '
        'Maryland, 17 U.S. 316 (1819).'
    )

    # REAL (correct) — McCulloch with its actual holding
    _add_para(doc,
        'The Necessary and Proper Clause grants Congress broad authority to select '
        'the means of executing its enumerated powers. McCulloch v. Maryland, '
        '17 U.S. (4 Wheat.) 316 (1819). The Court held that "the power to tax '
        'involves the power to destroy" and that states may not tax the operations '
        'of the federal government.'
    )

    _add_heading(doc, "V. WRONG YEAR FOR REAL CASE", level=2)

    # TRAP: Brown v. Board of Education with wrong year (1955 instead of 1954)
    _add_para(doc,
        'Separate educational facilities are inherently unequal and violate the '
        'Equal Protection Clause of the Fourteenth Amendment. Brown v. Board of '
        'Education, 347 U.S. 483 (1955).'
    )

    # TRAP: Gideon v. Wainwright with wrong year (1964 instead of 1963)
    _add_para(doc,
        'The Sixth Amendment\'s guarantee of the right to counsel is a '
        'fundamental right essential to a fair trial and is applicable to state '
        'criminal proceedings through the Fourteenth Amendment. Gideon v. '
        'Wainwright, 372 U.S. 335 (1964).'
    )

    _add_heading(doc, "VI. REAL CASES AS CONTROL", level=2)

    # REAL (correct) — Batson v. Kentucky
    _add_para(doc,
        'The Equal Protection Clause prohibits a prosecutor from using '
        'peremptory challenges to exclude jurors solely on the basis of race. '
        'Batson v. Kentucky, 476 U.S. 79, 89 (1986). A defendant may establish '
        'a prima facie case of purposeful discrimination by showing that the '
        'facts and circumstances raise an inference that the prosecutor used '
        'peremptory challenges on account of race.'
    )

    # REAL (correct) — Strickland v. Washington
    _add_para(doc,
        'To prevail on a claim of ineffective assistance of counsel, a defendant '
        'must show (1) that counsel\'s performance was deficient and (2) that the '
        'deficient performance prejudiced the defense. Strickland v. Washington, '
        '466 U.S. 668, 687 (1984). The Court noted that "judicial scrutiny of '
        'counsel\'s performance must be highly deferential."'
    )

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'For the foregoing reasons, Plaintiff respectfully requests that the '
        'Court rule in its favor on all claims.'
    )

    path = os.path.join(OUTPUT_DIR, "brief_k_subtle_traps.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


# ─────────────────────────────────────────────────────────────────────────
# BRIEF L: NEAR-MISS & LOOKALIKE CASES
# Tests: fake cases with names almost identical to real ones, party name
#   swaps (plaintiff/defendant reversed), real case wrong court, state
#   vs. federal confusion, "United States v." disambiguation, "In re"
#   cases that look similar to real ones
# ─────────────────────────────────────────────────────────────────────────

def create_brief_l():
    doc = Document()

    _add_heading(doc, "MEMORANDUM OF LAW IN SUPPORT OF PRELIMINARY INJUNCTION")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Apex Technologies, Inc. v. Consolidated Data Services, LLC")
    run.bold = True

    _add_heading(doc, "I. NEAR-MISS CASE NAMES", level=2)

    # FAKE — "Miranda v. State of Arizona" (real is "Miranda v. Arizona")
    _add_para(doc,
        'Suspects must be informed of their constitutional rights prior to '
        'custodial interrogation. Miranda v. State of Arizona, 384 U.S. 436 '
        '(1966). This requirement applies to all law enforcement agencies '
        'regardless of jurisdiction.'
    )

    # REAL — actual Miranda cite for comparison (should verify)
    _add_para(doc,
        'The warnings required prior to custodial interrogation serve to protect '
        'the Fifth Amendment privilege against self-incrimination. Miranda v. '
        'Arizona, 384 U.S. 436 (1966). The Court held that "the prosecution may '
        'not use statements, whether exculpatory or inculpatory, stemming from '
        'custodial interrogation of the defendant unless it demonstrates the use '
        'of procedural safeguards effective to secure the privilege against '
        'self-incrimination."'
    )

    _add_heading(doc, "II. PARTY NAME REVERSAL", level=2)

    # TRAP: Reversed parties — real is "New York Times Co. v. Sullivan"
    _add_para(doc,
        'A public official may not recover damages for a defamatory falsehood '
        'relating to his official conduct unless he proves the statement was made '
        'with "actual malice." Sullivan v. New York Times Co., 376 U.S. 254 '
        '(1964).'
    )

    # TRAP: Reversed parties — real is "Youngstown Sheet & Tube Co. v. Sawyer"
    _add_para(doc,
        'The President does not have inherent constitutional authority to seize '
        'private property in the absence of explicit congressional authorization. '
        'Sawyer v. Youngstown Sheet & Tube Co., 343 U.S. 579 (1952). Justice '
        'Jackson\'s concurrence established the influential three-category '
        'framework for analyzing presidential power.'
    )

    _add_heading(doc, "III. UNITED STATES v. [COMMON NAME] DISAMBIGUATION", level=2)

    # REAL — United States v. Nixon (Watergate tapes)
    _add_para(doc,
        'Executive privilege is not absolute and must yield to the demonstrated, '
        'specific need for evidence in a criminal trial. United States v. Nixon, '
        '418 U.S. 683, 713 (1974). The Court ordered President Nixon to comply '
        'with the subpoena for the White House tape recordings.'
    )

    # REAL — Nixon v. Fitzgerald (presidential immunity — different Nixon case)
    _add_para(doc,
        'A former President is entitled to absolute immunity from damages '
        'liability for acts within the "outer perimeter" of his official '
        'responsibility. Nixon v. Fitzgerald, 457 U.S. 731 (1982).'
    )

    # FAKE — "United States v. Williams" with fabricated holding
    # (There are dozens of real "United States v. Williams" cases —
    #  this uses a fake volume/reporter to test disambiguation)
    _add_para(doc,
        'The government bears the burden of proving that electronically stored '
        'information was not altered after seizure. United States v. Williams, '
        '893 F.3d 1127 (9th Cir. 2018) (holding that the Fourth Amendment '
        'requires a chain-of-custody foundation for all digital evidence '
        'introduced at trial).'
    )

    _add_heading(doc, "IV. STATE vs. FEDERAL VERSION CONFUSION", level=2)

    # REAL — Palsgraf (New York Court of Appeals, state court)
    _add_para(doc,
        'Liability in negligence requires a duty owed to the particular plaintiff. '
        'Palsgraf v. Long Island R.R. Co., 248 N.Y. 339, 162 N.E. 99 (1928). '
        'Chief Judge Cardozo held that negligence in the air, so to speak, will '
        'not do; there must be negligence directed at the plaintiff or a class '
        'of which the plaintiff is a member.'
    )

    # FAKE — Palsgraf with federal reporter (there is no federal Palsgraf)
    _add_para(doc,
        'The foreseeability of harm to a particular plaintiff is the touchstone '
        'of the duty analysis. Palsgraf v. Long Island R.R. Co., 248 F.2d 339 '
        '(2d Cir. 1928) (holding that proximate cause analysis must focus on the '
        'specific risk that made the actor\'s conduct negligent).'
    )

    _add_heading(doc, "V. IN RE CASES WITH SIMILAR NAMES", level=2)

    # REAL — In re Gault (juvenile due process)
    _add_para(doc,
        'Juveniles in delinquency proceedings are entitled to fundamental due '
        'process protections, including the right to notice, the right to '
        'counsel, the privilege against self-incrimination, and the right to '
        'confront witnesses. In re Gault, 387 U.S. 1 (1967).'
    )

    # FAKE — "In re Galt" (one letter off from Gault)
    _add_para(doc,
        'The Bankruptcy Code\'s automatic stay provisions apply to all '
        'proceedings against the debtor, including state regulatory actions '
        'seeking monetary penalties. In re Galt, 387 F.3d 295 (4th Cir. 2004) '
        '(distinguishing between governmental regulatory actions exempt under '
        'Section 362(b)(4) and those that are stayed).'
    )

    _add_heading(doc, "VI. MULTIPLE ISSUES IN SINGLE CITATION", level=2)

    # REAL — Jacobson v. Massachusetts (vaccine mandate, accurate)
    _add_para(doc,
        'States possess broad police power to enact reasonable regulations to '
        'protect public health, including compulsory vaccination laws. Jacobson '
        'v. Massachusetts, 197 U.S. 11 (1905). The Court upheld a Cambridge, '
        'Massachusetts ordinance requiring vaccination against smallpox, holding '
        'that individual liberty is not absolute and is subject to the restraint '
        'necessary for the common welfare.'
    )

    # TRAP: Correct case, fabricated modern extension
    _add_para(doc,
        'The Court extended Jacobson to hold that the government may mandate any '
        'medical procedure it deems necessary for public health without providing '
        'individualized medical exemptions. Jacobson v. Massachusetts, 197 U.S. '
        '11, 38 (1905). This unlimited public health authority has been '
        'consistently reaffirmed.'
    )

    _add_heading(doc, "CONCLUSION", level=2)

    _add_para(doc,
        'For the foregoing reasons, Plaintiff respectfully requests that the '
        'Court grant the preliminary injunction.'
    )

    path = os.path.join(OUTPUT_DIR, "brief_l_lookalikes.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    create_brief_a()
    create_brief_b()
    create_brief_c()
    create_brief_d()
    create_brief_e()
    create_brief_f()
    create_brief_g()
    create_brief_h()
    create_brief_i()
    create_brief_j()
    create_brief_k()
    create_brief_l()
    print("\nAll test briefs created successfully.")
